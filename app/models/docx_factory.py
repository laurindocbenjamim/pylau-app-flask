

import os

from docx import Document
from reportlab.pdfgen import canvas

class DocxFileFactory(object):

    # Function to save the response to a DOCX file
    def save_to_docx(*,content: str, filepath: str, filename: str="response.docx"):
        file_path=os.path.join(filepath,filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(file_path)
        return file_path

    # Save Optimised CV
    def save_as_docx(self, *,optimized_cv,file_path, filename):
        if optimized_cv=='' or file_path=='' or filename=='':
            return "File path, filename and content are required"

        doc = Document()
        doc.add_heading('Optimized CV', level=1)
        # Split markdown content into sections and add to Word doc
        for line in optimized_cv.split('\n'):
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            else:
                doc.add_paragraph(line)
        new_file=os.path.join(file_path,f"{filename}.docx")
        doc.save(new_file)
        return new_file
    
    
    # 
    def docx_to_pdf(self,*,docx_path: str, pdf_path: str):
        """ For simple DOCX files (text-only, basic formatting):
            
            Limitations:
            Loses images, tables, and complex formatting.
        """
        if docx_path=='' or pdf_path=='':
            return "Docx and PDF files path is required."
        
        try:
            doc = Document(docx_path)
            c = canvas.Canvas(pdf_path)

            y = 750  # Starting Y position
            for para in doc.paragraphs:
                c.drawString(50, y, para.text)
                y -= 15  # Adjust spacing

            c.save()
            True, pdf_path
        except Exception as e:
            return False, str(e)
    
    